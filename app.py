import telebot
from docx import Document
from docx.shared import Pt
import os

# --- Configuration ---
TELEGRAM_BOT_TOKEN = 'YOUR_TELEGRAM_BOT_TOKEN'
bot = telebot.TeleBot(TELEGRAM_BOT_TOKEN)

user_documents = {}

@bot.message_handler(commands=['start', 'help', 'new'])
def start_new_doc(message):
    chat_id = message.chat.id
    user_documents[chat_id] = []
    bot.reply_to(message, "📝 **Smart Document Started!**\n\nPaste your plain text here. I will automatically try to detect your Headings and Paragraphs without needing any special symbols. Send `/finish` when done.")

@bot.message_handler(commands=['finish'])
def compile_document(message):
    chat_id = message.chat.id
    
    if chat_id not in user_documents or not user_documents[chat_id]:
        bot.reply_to(message, "You haven't sent any text yet! Paste some text first.")
        return

    bot.reply_to(message, "Analyzing text structure and generating Word Document...")
    full_text = "\n\n".join(user_documents[chat_id])
    
    try:
        doc = Document()

        # Set the default global font to Arial, 12pt
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(12)

        lines = full_text.split('\n')
        
        for line in lines:
            clean_line = line.strip()
            if not clean_line:
                continue 
            
            # 1. Fallback for Markdown (If symbols accidentally survive)
            if clean_line.startswith('# '):
                doc.add_heading(clean_line[2:].strip(), level=1)
            elif clean_line.startswith('## ') or clean_line.startswith('### '):
                doc.add_heading(clean_line.lstrip('#').strip(), level=2)
            
            # 2. Detect standard phone/keyboard bullet points
            elif clean_line.startswith('•') or clean_line.startswith('- ') or clean_line.startswith('* '):
                # Remove the bullet character and add as a Word bullet list
                clean_text = clean_line.lstrip('•-* ').strip()
                doc.add_paragraph(clean_text, style='List Bullet')
            
            # 3. SMART AUTO-DETECT FOR HEADINGS (The Magic Fix)
            # If the line is short (under 65 chars) AND doesn't end with a punctuation mark, it's likely a title/heading!
            elif len(clean_line) < 65 and clean_line[-1] not in ".?!:,;":
                doc.add_heading(clean_line, level=1)
            
            # 4. Standard Paragraph
            else:
                doc.add_paragraph(clean_line)

        # Save the file
        filename = f"Smart_Project_{message.id}.docx"
        doc.save(filename)

        # Send it back
        with open(filename, 'rb') as doc_file:
            bot.send_document(chat_id, doc_file)

        if os.path.exists(filename):
            os.remove(filename)

        user_documents[chat_id] = []

    except Exception as e:
        bot.reply_to(message, f"An error occurred: {e}")

@bot.message_handler(func=lambda message: True)
def collect_text(message):
    chat_id = message.chat.id
    if chat_id not in user_documents:
        user_documents[chat_id] = []
    
    user_documents[chat_id].append(message.text)
    bot.reply_to(message, f"✅ Part {len(user_documents[chat_id])} saved! Send more, or send `/finish`.")

print("Smart Bot is running!")
bot.infinity_polling()
